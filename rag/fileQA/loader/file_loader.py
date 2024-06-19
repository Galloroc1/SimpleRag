import os.path
import sys
import traceback
import chardet
from abc import ABC, abstractmethod
from rag.fileQA.base import Document, MetaData
import logging
from bs4 import BeautifulSoup
import requests
import docx

logger = logging.getLogger(__name__)
import pymupdf


def detect_encoding(path):
    """
    check encoding
    :return:
    """
    with open(path, 'rb') as f:
        raw_data = f.read()
        result = chardet.detect(raw_data)
        encoding = result['encoding']
    return encoding


class BaseDataLoader(ABC):
    path = None

    @abstractmethod
    def load(self):
        raise

    def parse(self):
        raise


class PDFLoader(BaseDataLoader, ABC):

    def __init__(self, path):
        self.path = path

    def save_load_page(self, save_path):
        os.makedirs(save_path, exist_ok=True)
        pdf_document = pymupdf.open(self.path)
        for page in pdf_document:
            pix = page.get_pixmap()
            pix.save(os.path.join(save_path, f"page-{page.number}.png"))

    def load(self):
        text = ""
        try:
            full_text = []
            pdf_document = pymupdf.open(self.path)
            num_pages = pdf_document.page_count
            for page_num in range(num_pages):
                page = pdf_document.load_page(page_num)
                page_text = page.get_text(option='text')
                full_text.append(page_text)
            text = '\n'.join(full_text)
        except:
            traceback.print_exc()
            logger.debug("file read fail, we will return empty file", self.path)
        if len(text) == 0:
            raise f"{self.path} has not content"
        data = MetaData(meta=text, source={"path": self.path, "type": "pdf"})
        return Document([data], source={"path": self.path, "type": "pdf"})


class PDFLoaderUnstructured(BaseDataLoader, ABC):

    def __init__(self, path):
        self.path = path

    def load(self):
        pass


class CSVLoader(BaseDataLoader, ABC):

    def __init__(self, path):
        self.path = path

    def load(self):
        pass


class DocLoader(BaseDataLoader, ABC):

    def __init__(self, path):
        self.path = path

    def load(self):
        pass


class DocxLoader(BaseDataLoader, ABC):

    def __init__(self, path):
        self.path = path

    def load(self):
        doc = docx.Document(self.path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)

        data = MetaData(meta=text, source={"path": self.path, "type": "docx"})
        return Document([data], source={"path": self.path, "type": "docx"})


class ExcelLoader(BaseDataLoader, ABC):

    def __init__(self, path):
        self.path = path


class JsonLoader(BaseDataLoader, ABC):

    def __init__(self, path):
        self.path = path


class TxtLoader(BaseDataLoader, ABC):
    def __init__(self, path):
        self.path = path

    def load(self):
        text = ""
        try:
            with open(self.path, 'r', encoding=detect_encoding(self.path)) as f:
                text = f.read()
        except:
            traceback.print_exc()
            logger.debug("file read fail, we will return empty file", self.path)
        if len(text) == 0:
            raise f"{self.path} has not content"
        data = MetaData(meta=text, source={"path": self.path, "type": "txt"})
        return Document([data], source={"path": self.path, "type": "txt"})


class HtmlLoader(BaseDataLoader, ABC):

    def __init__(self, url, short_lens=20):
        self.path = url
        self.short_lens = short_lens

    def load(self):
        text = ""
        response = requests.get(self.path)
        if response.status_code == 200:
            try:
                encoding = response.apparent_encoding
                if encoding:
                    response.encoding = encoding
                else:
                    response.encoding = 'utf-8'  # 或者你可以使用一个默认值

                html_content = response.text
                soup = BeautifulSoup(html_content, 'html.parser')
                paragraphs = soup.find_all('p')

                for i, paragraph in enumerate(paragraphs):
                    now_text = paragraph.get_text(strip=True)
                    if len(now_text) > self.short_lens:
                        text = text + now_text + "\n"
            except:
                print("warning:has no sparse anything!")
        else:
            raise print(f'Failed to retrieve the page. Status code: {response.status_code}')
        data = MetaData(meta=text, source={"path": self.path, "type": "url"})
        return Document([data], source={"path": self.path, "type": "url"})
